"""
Extract and normalize commentary_extra documents into JSON chunks
for import into constitution-rag PostgreSQL.

Groups:
  1. Target audience docs (8 files) → commentary_extra_audiences_ru
  2. Referendum lines + counter-arguments + NK lines (3 files) → commentary_extra_lines_ru
  3. Key aspects + Commission theses (2 files) → commentary_extra_theses_ru
  4. Constitution comparison table (1 file) → commentary_extra_comparison_ru

Output: normalized/*.json files matching the existing schema.
"""

import json
import re
import unicodedata
from pathlib import Path
from docx import Document

RAW_DIR = Path("raw/commentary_extra")
OUT_DIR = Path("normalized")
OUT_DIR.mkdir(exist_ok=True)

LAYER = "commentary"
LANGUAGE = "ru"
EFFECTIVE_DATE = "2026-03-15"

# ── helpers ──────────────────────────────────────────────────────

def nfd(s: str) -> str:
    """Normalize to NFD to match filesystem."""
    return unicodedata.normalize("NFD", s)


def resolve_path(directory: Path, filename: str) -> Path:
    """Find a file matching filename with any Unicode normalization."""
    direct = directory / filename
    if direct.exists():
        return direct
    # Try NFD
    nfd_path = directory / nfd(filename)
    if nfd_path.exists():
        return nfd_path
    # Try NFC
    nfc_path = directory / unicodedata.normalize("NFC", filename)
    if nfc_path.exists():
        return nfc_path
    # Fallback: scan directory
    target = nfd(filename)
    for p in directory.iterdir():
        if nfd(p.name) == target:
            return p
    return direct  # return original (will fail with FileNotFoundError)


def clean(text: str) -> str:
    """Normalize whitespace and dashes."""
    text = text.strip()
    text = text.replace("\u00a0", " ")  # non-breaking space
    text = text.replace("\u2013", "–")  # en-dash
    text = text.replace("\u2014", "—")  # em-dash
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def paras_text(doc: Document) -> list[str]:
    """Return non-empty paragraph texts."""
    return [clean(p.text) for p in doc.paragraphs if p.text.strip()]


def make_chunk(
    chunk_id: str,
    heading: str,
    text: str,
    source_file: str,
    *,
    sub_group: str = "",
    audience: str = "",
    seq: int = 0,
) -> dict:
    return {
        "id": chunk_id,
        "layer": LAYER,
        "document": "commentary_extra",
        "language": LANGUAGE,
        "effective_date": EFFECTIVE_DATE,
        "source_file": source_file,
        "sub_group": sub_group,
        "audience": audience,
        "heading": heading,
        "text": text,
        "seq": seq,
        "status": "active",
    }


# ── Group 1: Target Audiences ───────────────────────────────────

AUDIENCE_FILES = {
    "Государственные_служащие": "gossluzh",
    "Журналисты": "zhurnalisty",
    "Молодёжь": "molodezh",
    "Правозащитники_и_активисты": "pravozashch",
    "Представители_бизнеса": "biznes",
    "Преподаватели": "prepodavateli",
    "Религиозные_служащие_и_представители": "religioznye",
    "национал_патриоты": "natpatrioty",
}


def extract_audiences() -> list[dict]:
    chunks = []
    for label, code in AUDIENCE_FILES.items():
        pattern = f"*{label}*"
        # Find file by pattern, handling Unicode normalization
        matches = [f for f in RAW_DIR.glob("*.docx") if label in f.name or nfd(label) in nfd(f.name)]
        if not matches:
            print(f"  WARN: no file for audience '{label}'")
            continue

        fpath = max(matches, key=lambda f: f.stat().st_size)  # pick non-empty
        doc = Document(str(fpath))
        texts = paras_text(doc)

        # First paragraph is always "Линии комментирования..."
        # Second is the audience heading (Heading 2)
        # Rest are the thesis paragraphs
        audience_name = ""
        thesis_texts = []
        for i, t in enumerate(texts):
            if i == 0 and "Линии комментирования" in t:
                continue
            if i <= 1 and not audience_name:
                audience_name = t
                continue
            thesis_texts.append(t)

        # Group short consecutive paragraphs into logical blocks
        # (some docs have multi-paragraph theses)
        merged = merge_short_paragraphs(thesis_texts, min_len=200)

        for seq, text in enumerate(merged, 1):
            chunk_id = f"ce_aud_{code}_{seq:02d}"
            heading_short = text[:80].split(".")[0].strip() if "." in text[:80] else text[:60].strip()
            chunks.append(make_chunk(
                chunk_id=chunk_id,
                heading=f"{audience_name}: {heading_short}",
                text=text,
                source_file=fpath.name,
                sub_group="target_audiences",
                audience=audience_name,
                seq=seq,
            ))
        print(f"  Audience '{audience_name}': {len(merged)} chunks from {fpath.name}")

    return chunks


def merge_short_paragraphs(texts: list[str], min_len: int = 200) -> list[str]:
    """Merge short paragraphs with the previous one to avoid micro-chunks."""
    if not texts:
        return []
    merged = [texts[0]]
    for t in texts[1:]:
        if len(merged[-1]) < min_len:
            merged[-1] = merged[-1] + "\n\n" + t
        else:
            merged.append(t)
    # Final pass: merge trailing short chunk
    if len(merged) > 1 and len(merged[-1]) < min_len:
        merged[-2] = merged[-2] + "\n\n" + merged[-1]
        merged.pop()
    return merged


# ── Group 2: Lines & Counter-arguments ──────────────────────────

def extract_lines() -> list[dict]:
    chunks = []

    # 2a: Линии Референдум — numbered theses (1., 2., 3., ...)
    fname = "26.02.11_Линии_Реферндум (1).docx"
    fpath = RAW_DIR / fname
    if fpath.exists():
        doc = Document(str(fpath))
        texts = paras_text(doc)
        # Group by numbered headers: lines starting with "N. "
        current_heading = ""
        current_body = []
        items = []

        for t in texts:
            # Skip top-level title
            if t in ("ЛИНИИ КОММЕНТИРОВАНИЯ", "Республиканский референдум по проекту новой Конституции"):
                continue
            # Check if this is a numbered heading
            m = re.match(r"^(\d+)\.\s+(.+)", t)
            if m and len(t) < 200:
                if current_heading:
                    items.append((current_heading, "\n\n".join(current_body)))
                current_heading = t
                current_body = []
            else:
                current_body.append(t)

        if current_heading:
            items.append((current_heading, "\n\n".join(current_body)))

        for seq, (heading, body) in enumerate(items, 1):
            chunk_id = f"ce_ref_line_{seq:02d}"
            chunks.append(make_chunk(
                chunk_id=chunk_id,
                heading=heading,
                text=body if body else heading,
                source_file=fname,
                sub_group="referendum_lines",
                seq=seq,
            ))
        print(f"  Referendum lines: {len(items)} chunks")

    # 2b: Контртезисы — numbered blocks grouped by topic
    fname = "26.02.12_ЛК_контртезисы_new (4).docx"
    fpath = RAW_DIR / fname
    if fpath.exists():
        doc = Document(str(fpath))
        texts = paras_text(doc)

        # Structure: top-level numbered topics (1. ЯЗЫКОВОЙ ВОПРОС, etc.)
        # then sub-numbered points under each
        current_topic = ""
        current_sub_heading = ""
        current_body = []
        items = []

        for t in texts:
            if t == "КОНТРАРГУМЕНТЫ ПО ОСНОВНЫМ ЛИНИЯМ КРИТИКИ":
                continue
            # Top-level topic: "1. ЯЗЫКОВОЙ ВОПРОС" (all caps or short numbered)
            m_topic = re.match(r"^(\d+)\.\s+([А-ЯЁ\s]+)$", t.strip())
            if m_topic:
                if current_topic and current_body:
                    items.append((current_topic, current_sub_heading, "\n\n".join(current_body)))
                elif current_sub_heading and current_body:
                    items.append((current_topic, current_sub_heading, "\n\n".join(current_body)))
                current_topic = t
                current_sub_heading = ""
                current_body = []
                continue

            # Sub-heading: "1. Исключение из новой Конституции..."
            m_sub = re.match(r"^(\d+)\.\s+(.{20,})", t)
            if m_sub and len(t) < 300 and not t[0].islower():
                if (current_sub_heading or current_topic) and current_body:
                    items.append((current_topic, current_sub_heading or current_topic, "\n\n".join(current_body)))
                current_sub_heading = t
                current_body = []
                continue

            current_body.append(t)

        if (current_sub_heading or current_topic) and current_body:
            items.append((current_topic, current_sub_heading or current_topic, "\n\n".join(current_body)))

        for seq, (topic, heading, body) in enumerate(items, 1):
            chunk_id = f"ce_counter_{seq:02d}"
            full_heading = f"{topic} — {heading}" if topic != heading else heading
            chunks.append(make_chunk(
                chunk_id=chunk_id,
                heading=full_heading,
                text=body,
                source_file=fname,
                sub_group="counter_arguments",
                seq=seq,
            ))
        print(f"  Counter-arguments: {len(items)} chunks")

    # 2c: ЛК по НК — lines on new constitution, numbered by section
    fname = "26.02.12_ЛК_по_НК_new (финал) (4).docx"
    fpath = RAW_DIR / fname
    if fpath.exists():
        doc = Document(str(fpath))
        texts = paras_text(doc)

        current_section = ""
        current_body = []
        items = []

        for t in texts:
            if t == "Линии комментирования по проекту новой Конституции РК":
                continue
            # Section headers: "Преамбула Конституции", "Раздел I. ...", etc.
            # or short descriptive lines that are section breaks
            is_section = False
            if re.match(r"^Раздел\s+[IVXLC]+\.", t):
                is_section = True
            elif t in ("Преамбула Конституции",):
                is_section = True
            elif len(t) < 120 and t[0].isupper() and not re.match(r"^\d+\.", t):
                # Short capitalized line — possible section break
                # But only if it looks like a heading
                if any(kw in t for kw in ("Раздел", "Преамбула", "Заключительные", "Переходные")):
                    is_section = True

            if is_section:
                if current_section and current_body:
                    items.append((current_section, "\n\n".join(current_body)))
                current_section = t
                current_body = []
            else:
                current_body.append(t)

        if current_section and current_body:
            items.append((current_section, "\n\n".join(current_body)))

        # If items are very large, split by numbered points within
        final_items = []
        for section, body in items:
            # Split by numbered points "1. ...", "2. ...", etc.
            parts = re.split(r"\n\n(?=\d+\.\s)", body)
            if len(parts) > 1:
                for part in parts:
                    m = re.match(r"^(\d+)\.\s+(.{10,})", part)
                    heading = f"{section}: п.{m.group(1)}" if m else section
                    final_items.append((heading, part))
            else:
                final_items.append((section, body))

        for seq, (heading, body) in enumerate(final_items, 1):
            chunk_id = f"ce_nk_line_{seq:02d}"
            chunks.append(make_chunk(
                chunk_id=chunk_id,
                heading=heading,
                text=body,
                source_file=fname,
                sub_group="nk_commentary_lines",
                seq=seq,
            ))
        print(f"  NK commentary lines: {len(final_items)} chunks")

    return chunks


# ── Group 3: Key Aspects + Commission Theses ────────────────────

def extract_theses() -> list[dict]:
    chunks = []

    # 3a: Ключевые аспекты
    fname = "Ключевые_аспекты_проекта_новой_Конституции.docx"
    fpath = resolve_path(RAW_DIR, fname)
    if fpath.exists():
        doc = Document(str(fpath))
        texts = paras_text(doc)

        # Structure: title + intro + numbered list paragraphs
        title = texts[0] if texts else ""
        items = []
        current = []

        for t in texts[1:]:
            # Each "List Paragraph" item starts a new point
            # They tend to start with capitalized descriptive phrase + "."
            if len(t) > 100:
                if current:
                    items.append("\n\n".join(current))
                current = [t]
            else:
                current.append(t)
        if current:
            items.append("\n\n".join(current))

        for seq, body in enumerate(items, 1):
            chunk_id = f"ce_aspects_{seq:02d}"
            heading_short = body[:80].split(".")[0].strip()
            chunks.append(make_chunk(
                chunk_id=chunk_id,
                heading=f"Ключевые аспекты: {heading_short}",
                text=body,
                source_file=fname,
                sub_group="key_aspects",
                seq=seq,
            ))
        print(f"  Key aspects: {len(items)} chunks")

    # 3b: Тезисы для Конституционной комиссии
    fname = "Тезисы_для_Конституционной_комиссии_1.docx"
    fpath = resolve_path(RAW_DIR, fname)
    if fpath.exists():
        doc = Document(str(fpath))
        texts = paras_text(doc)

        # Each paragraph is a standalone thesis
        title = texts[0] if texts else ""
        thesis_texts = texts[1:]  # skip title

        # Group consecutive short paragraphs
        merged = merge_short_paragraphs(thesis_texts, min_len=250)

        for seq, body in enumerate(merged, 1):
            chunk_id = f"ce_commission_{seq:02d}"
            heading_short = body[:80].split(".")[0].strip()
            chunks.append(make_chunk(
                chunk_id=chunk_id,
                heading=f"Тезисы Комиссии: {heading_short}",
                text=body,
                source_file=fname,
                sub_group="commission_theses",
                seq=seq,
            ))
        print(f"  Commission theses: {len(merged)} chunks")

    return chunks


# ── Group 4: Comparison Table ───────────────────────────────────

def extract_comparison() -> list[dict]:
    fname = "01.02.2026 - сравнение Конституций (2).docx"
    fpath = RAW_DIR / fname
    if not fpath.exists():
        print(f"  WARN: comparison file not found")
        return []

    doc = Document(str(fpath))
    if not doc.tables:
        print(f"  WARN: no table in comparison file")
        return []

    table = doc.tables[0]
    chunks = []

    # Headers: "Старая Конституция" | "Новая Конституция"
    # Rows come in pairs: topic row + content row
    # Group pairs into chunks

    rows_data = []
    for row in table.rows:
        cells = [clean(cell.text) for cell in row.cells]
        if all(c for c in cells):
            rows_data.append(cells)

    # Skip header row
    if rows_data and rows_data[0] == ["Старая Конституция", "Новая Конституция"]:
        rows_data = rows_data[1:]

    # Group by topic rows (short rows that match left/right = same topic)
    current_topic = ""
    current_pairs = []
    items = []

    for old, new in rows_data:
        if old == new and len(old) < 80:
            # This is a topic header
            if current_topic and current_pairs:
                items.append((current_topic, current_pairs))
            current_topic = old
            current_pairs = []
        else:
            current_pairs.append((old, new))

    if current_topic and current_pairs:
        items.append((current_topic, current_pairs))

    # If no topic grouping was found, treat each row as a chunk
    if not items:
        items = [("Сравнение", rows_data)]

    for seq, (topic, pairs) in enumerate(items, 1):
        text_parts = []
        for old, new in pairs:
            text_parts.append(f"Конституция 1995: {old}\nКонституция 2026: {new}")
        body = "\n\n".join(text_parts)

        chunk_id = f"ce_compare_{seq:02d}"
        chunks.append(make_chunk(
            chunk_id=chunk_id,
            heading=f"Сравнение: {topic}",
            text=body,
            source_file=fname,
            sub_group="comparison",
            seq=seq,
        ))

    print(f"  Comparison: {len(items)} chunks")
    return chunks


# ── Main ────────────────────────────────────────────────────────

def main():
    all_chunks = {}

    print("Extracting target audiences...")
    audiences = extract_audiences()
    all_chunks["commentary_extra_audiences_ru"] = audiences

    print("\nExtracting lines & counter-arguments...")
    lines = extract_lines()
    all_chunks["commentary_extra_lines_ru"] = lines

    print("\nExtracting theses...")
    theses = extract_theses()
    all_chunks["commentary_extra_theses_ru"] = theses

    print("\nExtracting comparison...")
    comparison = extract_comparison()
    all_chunks["commentary_extra_comparison_ru"] = comparison

    # Write JSON files
    for key, chunks in all_chunks.items():
        outpath = OUT_DIR / f"{key}_chunks.json"
        with open(outpath, "w", encoding="utf-8") as f:
            json.dump(chunks, f, ensure_ascii=False, indent=2)
        print(f"\n→ {outpath}: {len(chunks)} chunks, {outpath.stat().st_size / 1024:.1f} KB")

    total = sum(len(c) for c in all_chunks.values())
    print(f"\n═══ TOTAL: {total} chunks across {len(all_chunks)} datasets ═══")


if __name__ == "__main__":
    main()
